"""Module for processing Microsoft Word documents (.doc, .docx) to extract text content.

This module provides functionality to asynchronously and synchronously process Word
documents, extracting text from paragraphs and tables.

Example:
    Basic usage of the Word Processor:

    >>> processor = Processor()
    >>> text = await processor.process_file_async('document.docx', scanner)
    >>> print(text)  # Prints extracted text content
"""

import logging

from docx import Document


class Processor:
    """A processor class for handling Microsoft Word document text extraction.

    This class provides methods to extract text content from Word documents,
    including both paragraphs and tables. It supports both .doc and .docx formats
    and can process files both synchronously and asynchronously.

    Attributes:
        name (str): The display name of the processor.
        version (str): The version number of the processor.
        supported_extensions (list): List of supported file extensions.

    Example:
        >>> processor = Processor()
        >>> scanner = DocumentScanner()  # Your scanner instance
        >>> text = await processor.process_file_async('path/to/doc.docx', scanner)
        >>> print(text)
    """

    name = "Word Processor"
    version = "1.2"
    supported_extensions = [".doc", ".docx"]

    def __init__(self):
        """Initialize the Word Processor with a configured logger."""
        self.logger = logging.getLogger("processor-word")

    def __repr__(self):
        """Return a developer-friendly string representation of the processor.

        Returns:
            str: A string containing the class name, processor name, and version.
        """
        return f"<{self.__class__.__name__} {self.name} v{self.version}>"

    def __str__(self):
        """Return a user-friendly string representation of the processor.

        Returns:
            str: A string containing the processor name and version.
        """
        return f"{self.name} v{self.version}"

    async def process_file_async(self, file_path, scanner):
        """Process a Word document asynchronously.

        This method runs the synchronous processing in an executor to avoid
        blocking the event loop.

        Args:
            file_path (str): Path to the Word document to process.
            scanner: Scanner instance providing the event loop and executor.

        Returns:
            str: Extracted text content from the document. Empty string if processing fails.

        Example:
            >>> processor = Processor()
            >>> scanner = DocumentScanner()
            >>> text = await processor.process_file_async('report.docx', scanner)
        """
        try:
            loop = scanner.loop
            text_content = await loop.run_in_executor(
                scanner.executor, self.process_file, file_path
            )
            return text_content
        except Exception as e:
            self.logger.error(f"error processing: {file_path}: {e}")
            return ""

    def process_file(self, file_path):
        """Process a Word document synchronously.

        Extracts text content from both paragraphs and tables in the document.

        Args:
            file_path (str): Path to the Word document to process.

        Returns:
            str: Extracted text content from the document. Empty string if processing fails.

        Example:
            >>> processor = Processor()
            >>> text = processor.process_file('report.docx')
            >>> print(len(text))  # Print length of extracted text
        """
        text_content = ""
        try:
            self.logger.debug(f"opening Word document: {file_path}")
            doc = Document(file_path)
            for para_num, para in enumerate(doc.paragraphs, start=1):
                text_content += para.text + "\n"
                self.logger.debug(f"extracted text from paragraph: {para_num}")
            for table_num, table in enumerate(doc.tables, start=1):
                self.logger.debug(f"processing table: {table_num}")
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\n"
                        self.logger.debug("extracted text from table cell")
            return text_content
        except Exception as e:
            self.logger.error(f"error processing: {file_path}: {e}")
            return ""
